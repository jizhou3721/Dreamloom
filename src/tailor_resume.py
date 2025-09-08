# tailor_resume.py
"""
This is your original script, refactored to be importable.
Now it can be used by CLI OR Streamlit.
"""
from openai import OpenAI
client = OpenAI()

import os
import json
import argparse
import re
from typing import List, Tuple, Dict, Any

from docx import Document
from docx.shared import Pt, Inches

# -------------------------------
# CONFIGURATION
# -------------------------------
ATS_SAFE_FONT = "Calibri"
SECTION_HEADERS = ["EXPERIENCE", "EDUCATION", "TECH STACK", "CERTIFICATES"]
DEFAULT_FONT_SIZE_PT = 11.0
DEFAULT_LINE_HEIGHT_PT = 14.0

# -----------------------
# Helper utilities (unchanged)
# -----------------------
def get_pt_value(length_obj):
    try:
        return length_obj.pt if length_obj is not None else None
    except Exception:
        return None

def safe_style_name(style):
    try:
        return style.name if style is not None else ""
    except Exception:
        return ""

def is_bullet_paragraph(para):
    try:
        pPr = para._p.pPr
        if pPr is not None and getattr(pPr, 'numPr', None) is not None:
            return True
    except Exception:
        pass
    text = (para.text or "").strip()
    if text and text[:1] in ("→", "-", "•", "*"):
        return True
    style_name = safe_style_name(para.style).lower()
    if "bullet" in style_name or "list" in style_name or style_name == "list paragraph":
        return True
    return False

def clear_paragraph_runs(para):
    try:
        for run in list(para.runs):
            try:
                para._p.remove(run._r)
            except Exception:
                run.text = ""
    except Exception:
        try:
            para.text = ""
        except Exception:
            pass

def sanitize_text(text: str) -> str:
    if text is None:
        return ""
    replacements = {
        '→': '->',
        '•': '*',
        '’': "'",
        '‘': "'",
        '“': '"',
        '”': '"',
        '–': '-',
        '—': '-',
        '\u00A0': ' ',
    }
    for old, new in replacements.items():
        text = text.replace(old, new)
    return text

# -----------------------
# Extraction
# -----------------------
def extract_content(doc: Document) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
    content_blocks = []
    table_data = []

    for table in doc.tables:
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, cell_para in enumerate(cell.paragraphs):
                    runs_data = []
                    for run in cell_para.runs:
                        runs_data.append({
                            "text": run.text or "",
                            "bold": bool(run.bold),
                            "italic": bool(run.italic),
                            "font_size_pt": get_pt_value(run.font.size) or DEFAULT_FONT_SIZE_PT,
                            "font_name": run.font.name or ATS_SAFE_FONT,
                            "orig_length": len(run.text or "")
                        })
                    table_data.append({
                        "table": table,
                        "row_idx": r_idx,
                        "col_idx": c_idx,
                        "para_idx": p_idx,
                        "runs": runs_data,
                        "alignment": cell_para.alignment,
                        "style": safe_style_name(cell_para.style),
                        "text": cell_para.text or "",
                    })

    upper_headers = [h.upper() for h in SECTION_HEADERS]
    for para in doc.paragraphs:
        runs_data = []
        for run in para.runs:
            runs_data.append({
                "text": run.text or "",
                "bold": bool(run.bold),
                "italic": bool(run.italic),
                "font_size_pt": get_pt_value(run.font.size) or DEFAULT_FONT_SIZE_PT,
                "font_name": run.font.name or ATS_SAFE_FONT,
                "orig_length": len(run.text or "")
            })
        pf = para.paragraph_format
        content_blocks.append({
            "runs": runs_data,
            "style": safe_style_name(para.style),
            "alignment": para.alignment,
            "space_before_pt": get_pt_value(pf.space_before),
            "space_after_pt": get_pt_value(pf.space_after),
            "left_indent_pt": get_pt_value(pf.left_indent),
            "first_line_indent_pt": get_pt_value(pf.first_line_indent),
            "is_bullet": is_bullet_paragraph(para),
            "is_header": (para.text or "").strip().upper() in upper_headers,
            "text": para.text or ""
        })

    return content_blocks, table_data

def extract_full_resume_text(doc: Document) -> str:
    """
    Convert the resume DOCX into a plain-text version (for LLM input).
    Keeps section headers, bullet markers, and paragraph breaks.
    """
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        parts.append(text)
    # Also include table cells (if any)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    text = para.text.strip()
                    if text:
                        parts.append(text)
    return "\n".join(parts)


def generate_cover_letter(resume_path: str, job_description: str, output_docx: str):
    """
    Generate a professional cover letter tailored to the job description and resume.
    Saves result as a DOCX file.
    """
    doc = Document(resume_path)
    resume_text = extract_full_resume_text(doc)

    system_prompt = "You are an expert career coach and professional resume writer."
    user_prompt = f"""
Using the candidate's resume and the following job description, write a professional one-page cover letter.
- Tone: formal, confident, and aligned with industry best practices.
- Structure: introduction, skills alignment, closing with enthusiasm.
- Ensure it feels personalized and highlights the strongest overlaps.
- Auto fill the candidate's personal info based on the contents in the resume

Job description:
{job_description}

Candidate resume:
{resume_text}
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",  # use your preferred model
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.7,
    )

    cover_letter_text = response.choices[0].message.content.strip()


    # Save as DOCX
    out_doc = Document()
    for line in cover_letter_text.split("\n"):
        out_doc.add_paragraph(line)
    out_doc.save(output_docx)

    print(f"✅ Cover letter saved to: {output_docx}")


def generate_recruiter_message(resume_path: str, job_description: str, output_txt: str):
    """
    Generate a short LinkedIn recruiter message tailored to the job description and resume.
    Saves result as a TXT file.
    """
    doc = Document(resume_path)
    resume_text = extract_full_resume_text(doc)

    system_prompt = "You are an expert in professional networking and job search communication."
    user_prompt = f"""
Using the candidate's resume and the following job description, draft a short LinkedIn message
to a recruiter or hiring manager. Requirements:
- 2–3 sentences
- Polite, concise, and professional
- Show enthusiasm for the role
- Highlight one or two most relevant strengths
- Tell the recruiter the candidate's application is finished online
- Auto fill the candidate's personal info as the signature at the end of the drafte msg, based on the contents in the resume

Job description:
{job_description}

Candidate resume:
{resume_text}
"""

    response = client.chat.completions.create(
        model="gpt-4o-mini",  # use your preferred model
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": user_prompt},
        ],
        temperature=0.7,
    )

    recruiter_message = response.choices[0].message.content.strip()


    # Save as TXT
    with open(output_txt, "w", encoding="utf-8") as f:
        f.write(recruiter_message)
    
    print(f"✅ Recruiter message saved to: {output_txt}")


# -----------------------
# Build LLM input lines + line map
# -----------------------
def prepare_llm_input(content_blocks, table_data):
    lines = []
    line_map = []

    for i, block in enumerate(content_blocks):
        prefix_parts = []
        if block["is_header"]:
            prefix_parts.append("[SECTION_HEADER]")
        if block["is_bullet"]:
            prefix_parts.append("[BULLET]")
        if block.get("style"):
            prefix_parts.append(f"[STYLE={block['style']}]")
        if block.get("alignment") is not None:
            prefix_parts.append(f"[ALIGN={str(block['alignment'])}]")
        prefix = " ".join(prefix_parts)
        line = (prefix + " " if prefix else "") + (block.get("text") or "")
        lines.append(line)
        line_map.append({"kind": "para", "para_idx": i})

    for t_idx, td in enumerate(table_data):
        prefix_parts = ["[TABLE_CELL]"]
        if td.get("style"):
            prefix_parts.append(f"[STYLE={td['style']}]")
        if td.get("alignment") is not None:
            prefix_parts.append(f"[ALIGN={str(td['alignment'])}]")
        prefix = " ".join(prefix_parts)
        line = prefix + " " + (td.get("text") or "")
        lines.append(line)
        line_map.append({"kind": "table", "table_idx": t_idx})

    return lines, line_map

# -----------------------
# LLM call (robust JSON)
# -----------------------
def llm_call(lines: List[str], job_description: str) -> List[str]:
    n = len(lines)
    system_rules = (
        "You are a resume tailoring assistant. "
        "You must rewrite the resume line-by-line to best match a job description. "
        "Follow these rewrite rules for each bullet or key point: \n"
        "1. If the current line is NOT relevant enough to the job description, "
        "   completely rewrite it into a more relevant achievement, "
        "   drawing on related experience from elsewhere in the resume. "
        "   Do not invent new employers, roles, or dates, but you may adapt responsibilities/projects. \n"
        "2. If the current line IS relevant, expand it to emphasize stronger impact, "
        "   integrate job-specific skills/technologies, and highlight measurable outcomes. \n"
        "Always:\n"
        "- Keep a strict 1-to-1 mapping (one input line → one output line).\n"
        "- Preserve section headers and bullet markers.\n"
        "- Use strong action verbs and quantified impact when possible.\n"
        "- Ensure the entire output is a JSON array of strings, same length as input."
    )

    section_headers_str = ', '.join(f'"{h}"' for h in SECTION_HEADERS)

    user_prompt = f"""
You are tailoring the following resume to the job description below.

JOB DESCRIPTION (summarized or original):
{job_description}

Rules:
- Output a **JSON array of strings** (no code fences, no extra text).
- The JSON array must contain exactly {n} elements (same number of lines as input).
- Each element corresponds to the input line at the same index.
- Preserve structure: keep [SECTION_HEADER] unchanged, and keep bullet markers if present.
- You must actively rewrite content to emphasize relevant skills, technologies, and achievements.
- Do NOT just rephrase slightly — meaningfully integrate job-specific skills into each line if possible.
- When the line is unrelated to the job description (e.g., education details), leave it mostly unchanged.
- Do not invent new employers, roles, or dates, but you may reasonably expand responsibilities/projects
  with technologies from the job description.

Canonical section headers (keep verbatim): {section_headers_str}

Input resume lines (array of {n} items):
{json.dumps(lines, ensure_ascii=False)}
"""

    use_openai = os.getenv("USE_OPENAI", "0") == "1"
    raw = None
    if use_openai:
        api_key = os.getenv("OPENAI_API_KEY")
        if not api_key:
            raise RuntimeError("OPENAI_API_KEY not set while USE_OPENAI=1.")
        try:
            chat = client.chat.completions.create(
                model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
                messages=[
                    {"role": "system", "content": system_rules},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.3,
            )
            raw = chat.choices[0].message.content.strip()
        except Exception as e:
            raise RuntimeError(f"OpenAI call failed: {e}")
    else:
        print("\n========== LLM PROMPT (BEGIN) ==========")
        print(system_rules)
        print(user_prompt)
        print("=========== LLM PROMPT (END) ===========\n")
        raw = input("Paste JSON array output from LLM here:\n").strip()

    try:
        data = json.loads(raw)
    except Exception as e:
        m = re.search(r"(\[.*\])", raw, flags=re.S)
        if m:
            try:
                data = json.loads(m.group(1))
            except Exception as e2:
                raise ValueError(f"LLM output is not valid JSON: {e2}\nRaw: {raw[:500]}...")
        else:
            raise ValueError(f"LLM output is not valid JSON: {e}\nRaw: {raw[:500]}...")

    if not isinstance(data, list):
        raise ValueError("LLM output must be a JSON array.")

    out = []
    for elem in data:
        if elem is None:
            out.append("")
        elif isinstance(elem, (str, int, float)):
            out.append(str(elem))
        else:
            try:
                out.append(json.dumps(elem, ensure_ascii=False))
            except Exception:
                out.append("")

    if len(out) < n:
        out += [""] * (n - len(out))
    elif len(out) > n:
        out = out[:n]

    out = [sanitize_text(s) for s in out]
    return out

# -----------------------
# Reinjection
# -----------------------
def reinject_content(doc: Document, tailored_texts: List[str],
                     content_blocks: List[Dict[str, Any]],
                     table_data: List[Dict[str, Any]],
                     line_map: List[Dict[str, Any]]) -> None:
    def proportional_chunks(text: str, run_specs: List[Dict[str, Any]]) -> List[str]:
        if not run_specs:
            return [text]
        total_len = sum(r.get("orig_length", 0) or 0 for r in run_specs)
        n = len(run_specs)
        if total_len <= 0:
            size = len(text) // n
            sizes = [size] * n
            sizes[-1] += len(text) - sum(sizes)
        else:
            sizes = [int(round((r.get("orig_length", 0) or 0) / total_len * len(text))) for r in run_specs]
            diff = len(text) - sum(sizes)
            if diff != 0:
                sizes[-1] = max(0, sizes[-1] + diff)
        chunks, pos = [], 0
        for s in sizes:
            chunks.append(text[pos:pos + s] if s > 0 else "")
            pos += s
        return chunks

    for lm_idx, lm in enumerate(line_map):
        line_text = tailored_texts[lm_idx]
        if lm["kind"] == "para":
            i = lm["para_idx"]
            if i >= len(doc.paragraphs):
                continue
            para = doc.paragraphs[i]
            block = content_blocks[i]
            clear_paragraph_runs(para)

            chunks = proportional_chunks(line_text, block["runs"])
            for r_spec, chunk in zip(block["runs"], chunks):
                run = para.add_run(chunk)
                run.bold = bool(r_spec.get("bold"))
                run.italic = bool(r_spec.get("italic"))
                try:
                    run.font.size = Pt(float(r_spec.get("font_size_pt") or DEFAULT_FONT_SIZE_PT))
                except Exception:
                    pass
                try:
                    run.font.name = r_spec.get("font_name") or ATS_SAFE_FONT
                except Exception:
                    pass

            fmt = para.paragraph_format
            if block.get("space_before_pt") is not None:
                fmt.space_before = Pt(block["space_before_pt"])
            if block.get("space_after_pt") is not None:
                fmt.space_after = Pt(block["space_after_pt"])
            if block.get("left_indent_pt") is not None:
                fmt.left_indent = Pt(block["left_indent_pt"])
            if block.get("first_line_indent_pt") is not None:
                fmt.first_line_indent = Pt(block["first_line_indent_pt"])
            if block.get("alignment") is not None:
                para.alignment = block["alignment"]

            try:
                style_name = block.get("style") or "Normal"
                if style_name in doc.styles:
                    para.style = doc.styles[style_name]
            except Exception:
                pass

        elif lm["kind"] == "table":
            td_index = lm["table_idx"]
            if td_index >= len(table_data):
                continue
            td = table_data[td_index]
            table = td["table"]
            row = table.rows[td["row_idx"]]
            cell = row.cells[td["col_idx"]]
            p_idx = td.get("para_idx", 0)
            while len(cell.paragraphs) <= p_idx:
                cell.add_paragraph("")
            cell_para = cell.paragraphs[p_idx]

            clear_paragraph_runs(cell_para)
            chunks = proportional_chunks(line_text, td["runs"])
            for r_spec, chunk in zip(td["runs"], chunks):
                run = cell_para.add_run(chunk)
                run.bold = bool(r_spec.get("bold"))
                run.italic = bool(r_spec.get("italic"))
                try:
                    run.font.size = Pt(float(r_spec.get("font_size_pt") or DEFAULT_FONT_SIZE_PT))
                except Exception:
                    pass
                try:
                    run.font.name = r_spec.get("font_name") or ATS_SAFE_FONT
                except Exception:
                    pass

            if td.get("alignment") is not None:
                cell_para.alignment = td["alignment"]
            try:
                if td.get("style") and td["style"] in doc.styles:
                    cell_para.style = doc.styles[td["style"]]
            except Exception:
                pass

# -----------------------
# Summarize Job Description
# -----------------------
def summarize_job_description(job_description: str, max_points: int = 4) -> str:
    system_prompt = (
        "You are a helpful assistant that extracts the most relevant information "
        "from job descriptions for resume tailoring."
    )
    user_prompt = f"""
Please summarize the following job description into {max_points} concise bullet points.
Focus on:
- Required technical skills
- Key responsibilities
- Important soft skills or certifications
- Must-have qualifications

Job description:
{job_description}
"""

    use_openai = os.getenv("USE_OPENAI", "0") == "1"
    if use_openai:
        chat = client.chat.completions.create(
            model=os.getenv("OPENAI_MODEL", "gpt-4o-mini"),
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.3,
        )
        summary = chat.choices[0].message.content.strip()
    else:
        print("\n========== JD SUMMARIZATION PROMPT ==========")
        print(user_prompt)
        print("=============================================\n")
        summary = input("Paste JD summary here:\n").strip()
    return summary

# -----------------------
# Main Function (for Streamlit to call)
# -----------------------
def tailor_resume_in_memory(input_docx_path: str, job_description: str, enforce_one_page: bool = True) -> bytes:
    """
    Runs the full tailoring process and returns the tailored .docx as bytes.
    """
    from docx import Document

    # Step 1: Summarize JD if long
    jd_to_use = job_description
    if len(job_description.split()) > 150:
        jd_to_use = summarize_job_description(job_description)

    # Step 2: Load and process
    doc = Document(input_docx_path)
    content_blocks, table_data = extract_content(doc)
    lines, line_map = prepare_llm_input(content_blocks, table_data)
    tailored_lines = llm_call(lines, jd_to_use)

    # Ensure length match
    if len(tailored_lines) < len(lines):
        tailored_lines += [""] * (len(lines) - len(tailored_lines))
    elif len(tailored_lines) > len(lines):
        tailored_lines = tailored_lines[:len(lines)]

    # Step 3: Reinject
    reinject_content(doc, tailored_lines, content_blocks, table_data, line_map)

    if enforce_one_page:
        # Optional: add truncate logic here if needed
        pass

    # Save to bytes
    from io import BytesIO
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()

if __name__ == "__main__":
    import streamlit as st

    st.title("Resume Tailor Pro")

    # File uploader for resume
    uploaded_file = st.file_uploader("Upload your resume (.docx)", type=["docx"])

    # Text area for job description
    job_description = st.text_area("Job description")

    if st.button("Tailor Resume"):
        if not uploaded_file or not job_description.strip():
            st.error("Please upload a resume and provide a job description.")
        else:
            uploaded_resume_bytes = uploaded_file.read()
            with open("temp_resume.docx", "wb") as f:
                f.write(uploaded_resume_bytes)

            tailored_bytes = tailor_resume.tailor_resume_in_memory("temp_resume.docx", job_description)
            with open("resume_tailored.docx", "wb") as f:
                f.write(tailored_bytes)

            cover_letter_path = "resume_tailored_cover_letter.docx"
            tailor_resume.generate_cover_letter("temp_resume.docx", job_description, cover_letter_path)
            with open(cover_letter_path, "rb") as f:
                cover_letter_bytes = f.read()

            recruiter_msg_path = "resume_tailored_recruiter_msg.txt"
            tailor_resume.generate_recruiter_message("temp_resume.docx", job_description, recruiter_msg_path)
            with open(recruiter_msg_path, "rb") as f:
                recruiter_msg_bytes = f.read()

            # Store in session state
            st.session_state["tailored_bytes"] = tailored_bytes
            st.session_state["cover_letter_bytes"] = cover_letter_bytes
            st.session_state["recruiter_msg_bytes"] = recruiter_msg_bytes
            st.success("Resume tailored! Download your files below.")

    # Show download buttons if files exist in session state
    if "tailored_bytes" in st.session_state:
        st.download_button(
            "Download Tailored Resume (.docx)",
            st.session_state["tailored_bytes"],
            file_name="resume_tailored.docx"
        )
    if "cover_letter_bytes" in st.session_state:
        st.download_button(
            "Download Cover Letter (.docx)",
            st.session_state["cover_letter_bytes"],
            file_name="resume_tailored_cover_letter.docx"
        )
    if "recruiter_msg_bytes" in st.session_state:
        st.download_button(
            "Download Recruiter Message (.txt)",
            st.session_state["recruiter_msg_bytes"],
            file_name="resume_tailored_recruiter_msg.txt"
        )